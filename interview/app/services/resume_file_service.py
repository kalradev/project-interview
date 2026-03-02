"""Extract plain text from uploaded resume files (PDF, DOCX) for parsing."""

import logging
from io import BytesIO
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_resume_file(content: bytes, filename: str) -> Optional[str]:
    """
    Extract plain text from a resume file. Supports PDF and DOCX.
    Returns None if format is unsupported or extraction fails.
    """
    if not content or not filename:
        return None
    name_lower = filename.lower()
    if name_lower.endswith(".pdf"):
        return _extract_pdf(content)
    if name_lower.endswith(".docx"):
        return _extract_docx(content)
    return None


def _extract_pdf(content: bytes) -> Optional[str]:
    # Prefer pdfplumber for layout preservation (better for multi-column resumes)
    try:
        import pdfplumber
        with pdfplumber.open(BytesIO(content)) as pdf:
            parts = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
            result = "\n".join(parts).strip() or None
            if result:
                return result
    except Exception as e:
        logger.debug("pdfplumber PDF extraction failed, trying pypdf: %s", e)
    try:
        from pypdf import PdfReader
        reader = PdfReader(BytesIO(content))
        parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
        result = "\n".join(parts).strip() or None
        if not result:
            logger.warning("PDF extraction returned no text (file may be scanned/image-only)")
        return result
    except Exception as e:
        logger.warning("PDF extraction failed: %s", e)
        return None


def _extract_docx(content: bytes) -> Optional[str]:
    try:
        from docx import Document
        doc = Document(BytesIO(content))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        result = "\n".join(parts).strip() or None
        if not result:
            logger.warning("DOCX extraction returned no text")
        return result
    except Exception as e:
        logger.warning("DOCX extraction failed: %s", e)
        return None
